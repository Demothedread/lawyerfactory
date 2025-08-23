"""
# Script Name: legal_document_generator.py
# Description: Legal Document Generator for LawyerFactory Orchestration Phase.  This module handles the generation of professional legal documents in multiple formats: - Microsoft Word (.docx) with proper legal formatting - PDF with court-ready styling - Markdown for UI display and collaboration  Supports the complete legal document structure: - Statement of Facts - Claims of Action (with IRAC methodology) - Prayer for Relief - Remedies  Integrates with court filing requirements and Bluebook citation standards.
# Relationships:
#   - Entity Type: Module
#   - Directory Group: Document Generation
#   - Group Tags: null
Legal Document Generator for LawyerFactory Orchestration Phase.

This module handles the generation of professional legal documents in multiple formats:
- Microsoft Word (.docx) with proper legal formatting
- PDF with court-ready styling
- Markdown for UI display and collaboration

Supports the complete legal document structure:
- Statement of Facts
- Claims of Action (with IRAC methodology)
- Prayer for Relief
- Remedies

Integrates with court filing requirements and Bluebook citation standards.
"""

import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any
import json

logger = logging.getLogger(__name__)


class LegalDocumentGenerator:
    """Main controller for legal document generation"""

    def __init__(self, output_dir: str = "output/orchestration"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True, parents=True)

        # Initialize format generators
        self.format_generators = {
            'word': WordDocumentGenerator(),
            'pdf': PDFDocumentGenerator(),
            'markdown': MarkdownDocumentGenerator()
        }

    async def generate_document(self, document_type: str, content: Dict[str, Any],
                              format: str = 'word', metadata: Optional[Dict[str, Any]] = None) -> str:
        """Generate a legal document in the specified format"""

        try:
            if format not in self.format_generators:
                raise ValueError(f"Unsupported format: {format}")

            generator = self.format_generators[format]

            # Generate the document
            file_path = await generator.generate(document_type, content, metadata or {})

            logger.info(f"Generated {format} document: {file_path}")
            return str(file_path)

        except Exception as e:
            logger.error(f"Document generation failed: {e}")
            raise

    async def generate_all_formats(self, document_type: str, content: Dict[str, Any],
                                 metadata: Optional[Dict[str, Any]] = None) -> Dict[str, str]:
        """Generate the same document in all supported formats"""

        results = {}
        for format_name, generator in self.format_generators.items():
            try:
                file_path = await generator.generate(document_type, content, metadata or {})
                results[format_name] = str(file_path)
            except Exception as e:
                logger.error(f"Failed to generate {format_name} format: {e}")
                results[format_name] = f"Error: {str(e)}"

        return results

    async def generate_statement_of_facts(self, facts: List[str], metadata: Dict[str, Any]) -> Dict[str, str]:
        """Generate Statement of Facts document"""

        content = {
            'title': 'STATEMENT OF FACTS',
            'facts': facts,
            'jurisdiction': metadata.get('jurisdiction', 'California'),
            'court': metadata.get('court', 'Superior Court'),
            'case_number': metadata.get('case_number', 'CV-2024-XXXXX'),
            'parties': metadata.get('parties', [])
        }

        return await self.generate_all_formats('statement_of_facts', content, metadata)

    async def generate_claims_of_action(self, claims: List[Dict[str, Any]], metadata: Dict[str, Any]) -> Dict[str, str]:
        """Generate Claims of Action document with IRAC methodology"""

        content = {
            'title': 'CLAIMS OF ACTION',
            'claims': claims,
            'use_irac': True,
            'jurisdiction': metadata.get('jurisdiction', 'California'),
            'case_number': metadata.get('case_number', 'CV-2024-XXXXX')
        }

        return await self.generate_all_formats('claims_of_action', content, metadata)

    async def generate_prayer_for_relief(self, remedies: List[str], metadata: Dict[str, Any]) -> Dict[str, str]:
        """Generate Prayer for Relief document"""

        content = {
            'title': 'PRAYER FOR RELIEF',
            'remedies': remedies,
            'jurisdiction': metadata.get('jurisdiction', 'California'),
            'case_number': metadata.get('case_number', 'CV-2024-XXXXX')
        }

        return await self.generate_all_formats('prayer_for_relief', content, metadata)

    async def generate_remedies_section(self, damages: Dict[str, Any], metadata: Dict[str, Any]) -> Dict[str, str]:
        """Generate Remedies/Damages section"""

        content = {
            'title': 'REMEDIES',
            'damages': damages,
            'jurisdiction': metadata.get('jurisdiction', 'California'),
            'case_number': metadata.get('case_number', 'CV-2024-XXXXX')
        }

        return await self.generate_all_formats('remedies', content, metadata)


class WordDocumentGenerator:
    """Generates Microsoft Word documents with proper legal formatting"""

    def __init__(self):
        self.output_dir = Path("output/orchestration/word")
        self.output_dir.mkdir(exist_ok=True, parents=True)

    async def generate(self, document_type: str, content: Dict[str, Any], metadata: Dict[str, Any]) -> str:
        """Generate a Word document"""

        try:
            # Import here to avoid dependency issues
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE

            doc = Document()

            # Set up document properties
            doc.core_properties.title = content.get('title', 'Legal Document')
            doc.core_properties.author = 'LawyerFactory AI Maestro'

            # Add header
            header = doc.sections[0].header
            header_text = f"{metadata.get('case_name', 'Unknown Case')} - {content.get('title', 'Document')}"
            header.paragraphs[0].text = header_text

            # Generate content based on type
            if document_type == 'statement_of_facts':
                await self._generate_statement_of_facts(doc, content)
            elif document_type == 'claims_of_action':
                await self._generate_claims_of_action(doc, content)
            elif document_type == 'prayer_for_relief':
                await self._generate_prayer_for_relief(doc, content)
            elif document_type == 'remedies':
                await self._generate_remedies(doc, content)
            else:
                await self._generate_generic_document(doc, content)

            # Save document
            filename = f"{document_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            file_path = self.output_dir / filename
            doc.save(str(file_path))

            return str(file_path)

        except ImportError:
            logger.warning("python-docx not available, generating placeholder")
            return await self._generate_placeholder(document_type, 'docx')

    async def _generate_statement_of_facts(self, doc, content: Dict[str, Any]):
        """Generate Statement of Facts section"""

        # Title
        title = doc.add_heading(content['title'], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Introduction
        intro = doc.add_paragraph()
        intro.add_run(f"This action is brought in the {content['jurisdiction']} {content['court']}.").italic = True

        doc.add_paragraph()  # Spacing

        # Facts in chronological order
        facts = content.get('facts', [])
        for i, fact in enumerate(facts, 1):
            p = doc.add_paragraph(f"{i}. {fact}")
            p.style = 'Normal'

        doc.add_paragraph()  # Spacing

    async def _generate_claims_of_action(self, doc, content: Dict[str, Any]):
        """Generate Claims of Action with IRAC methodology"""

        title = doc.add_heading(content['title'], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        claims = content.get('claims', [])
        for i, claim in enumerate(claims, 1):
            if isinstance(claim, dict):
                claim_title = claim.get('title', f'Claim {i}')
                elements = claim.get('elements', [])
                irac_analysis = claim.get('irac_analysis', {})

                # Claim heading
                doc.add_heading(f'{i}. {claim_title}', level=1)

                # IRAC structure
                if content.get('use_irac', False) and irac_analysis:
                    await self._add_irac_section(doc, irac_analysis)
                else:
                    # Fallback to basic elements
                    for element in elements:
                        doc.add_paragraph(f"• {element}")

                doc.add_paragraph()  # Spacing between claims

    async def _add_irac_section(self, doc, irac_analysis: Dict[str, Any]):
        """Add IRAC methodology section"""

        # Issue
        issue_heading = doc.add_heading('ISSUE', level=2)
        issue_text = irac_analysis.get('issue', 'Whether...')
        doc.add_paragraph(issue_text)

        # Rule
        rule_heading = doc.add_heading('RULE', level=2)
        rule_text = irac_analysis.get('rule', 'The applicable rule is...')
        doc.add_paragraph(rule_text)

        # Application/Analysis
        analysis_heading = doc.add_heading('APPLICATION/ANALYSIS', level=2)
        analysis_text = irac_analysis.get('analysis', 'Applying the rule to the facts...')
        doc.add_paragraph(analysis_text)

        # Counterarguments
        counterargs = irac_analysis.get('counterarguments', [])
        if counterargs:
            counter_heading = doc.add_heading('COUNTERARGUMENTS', level=2)
            for counter in counterargs:
                doc.add_paragraph(f"• {counter}")

        # Conclusion
        conclusion_heading = doc.add_heading('CONCLUSION', level=2)
        conclusion_text = irac_analysis.get('conclusion', 'Therefore...')
        doc.add_paragraph(conclusion_text)

    async def _generate_prayer_for_relief(self, doc, content: Dict[str, Any]):
        """Generate Prayer for Relief section"""

        title = doc.add_heading(content['title'], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_heading('WHEREFORE, Plaintiff prays for judgment as follows:', level=1)

        remedies = content.get('remedies', [])
        for i, remedy in enumerate(remedies, 1):
            doc.add_paragraph(f"{i}. {remedy}")

        doc.add_paragraph()  # Spacing
        doc.add_paragraph("Respectfully submitted,")

    async def _generate_remedies(self, doc, content: Dict[str, Any]):
        """Generate Remedies section"""

        title = doc.add_heading(content['title'], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        damages = content.get('damages', {})

        if isinstance(damages, dict):
            for damage_type, amount in damages.items():
                doc.add_paragraph(f"{damage_type}: {amount}")

        doc.add_paragraph()  # Spacing

    async def _generate_generic_document(self, doc, content: Dict[str, Any]):
        """Generate a generic document"""

        title = doc.add_heading(content.get('title', 'Legal Document'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add content sections
        for key, value in content.items():
            if key != 'title' and isinstance(value, (str, list)):
                doc.add_heading(key.replace('_', ' ').title(), level=1)
                if isinstance(value, list):
                    for item in value:
                        doc.add_paragraph(f"• {item}")
                else:
                    doc.add_paragraph(value)

    async def _generate_placeholder(self, document_type: str, extension: str) -> str:
        """Generate a placeholder file when dependencies are missing"""

        filename = f"{document_type}_placeholder_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{extension}"
        file_path = self.output_dir / filename

        with open(file_path, 'w') as f:
            f.write(f"Placeholder {extension.upper()} document: {document_type}\n")
            f.write("Generated by LawyerFactory AI Maestro\n")
            f.write(f"Timestamp: {datetime.now().isoformat()}\n")

        return str(file_path)


class PDFDocumentGenerator:
    """Generates PDF documents with court-ready styling"""

    def __init__(self):
        self.output_dir = Path("output/orchestration/pdf")
        self.output_dir.mkdir(exist_ok=True, parents=True)

    async def generate(self, document_type: str, content: Dict[str, Any], metadata: Dict[str, Any]) -> str:
        """Generate a PDF document"""

        try:
            # Import here to avoid dependency issues
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.lib.units import inch
            from reportlab.lib.enums import TA_CENTER, TA_LEFT

            filename = f"{document_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            file_path = self.output_dir / filename

            doc = SimpleDocTemplate(str(file_path), pagesize=letter)
            styles = getSampleStyleSheet()
            story = []

            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                alignment=TA_CENTER,
                spaceAfter=30,
            )

            normal_style = styles['Normal']

            # Generate content based on type
            if document_type == 'statement_of_facts':
                await self._generate_statement_of_facts_pdf(story, content, title_style, normal_style)
            elif document_type == 'claims_of_action':
                await self._generate_claims_of_action_pdf(story, content, title_style, normal_style)
            elif document_type == 'prayer_for_relief':
                await self._generate_prayer_for_relief_pdf(story, content, title_style, normal_style)
            elif document_type == 'remedies':
                await self._generate_remedies_pdf(story, content, title_style, normal_style)
            else:
                await self._generate_generic_pdf(story, content, title_style, normal_style)

            doc.build(story)
            return str(file_path)

        except ImportError:
            logger.warning("reportlab not available, generating placeholder")
            return await self._generate_placeholder(document_type, 'pdf')

    async def _generate_statement_of_facts_pdf(self, story, content, title_style, normal_style):
        """Generate Statement of Facts PDF"""

        story.append(Paragraph(content['title'], title_style))
        story.append(Spacer(1, 20))

        intro_text = f"This action is brought in the {content['jurisdiction']} {content['court']}."
        story.append(Paragraph(f'<i>{intro_text}</i>', normal_style))
        story.append(Spacer(1, 20))

        facts = content.get('facts', [])
        for i, fact in enumerate(facts, 1):
            story.append(Paragraph(f"{i}. {fact}", normal_style))
            story.append(Spacer(1, 12))

    async def _generate_claims_of_action_pdf(self, story, content, title_style, normal_style):
        """Generate Claims of Action PDF with IRAC"""

        story.append(Paragraph(content['title'], title_style))
        story.append(Spacer(1, 20))

        claims = content.get('claims', [])
        for i, claim in enumerate(claims, 1):
            if isinstance(claim, dict):
                claim_title = claim.get('title', f'Claim {i}')
                story.append(Paragraph(f'{i}. {claim_title}', styles['Heading2']))
                story.append(Spacer(1, 12))

                if content.get('use_irac', False):
                    irac_analysis = claim.get('irac_analysis', {})
                    await self._add_irac_pdf(story, irac_analysis, styles)

                story.append(Spacer(1, 20))

    async def _add_irac_pdf(self, story, irac_analysis, styles):
        """Add IRAC methodology to PDF"""

        if irac_analysis.get('issue'):
            story.append(Paragraph('ISSUE', styles['Heading3']))
            story.append(Paragraph(irac_analysis['issue'], styles['Normal']))
            story.append(Spacer(1, 12))

        if irac_analysis.get('rule'):
            story.append(Paragraph('RULE', styles['Heading3']))
            story.append(Paragraph(irac_analysis['rule'], styles['Normal']))
            story.append(Spacer(1, 12))

        if irac_analysis.get('analysis'):
            story.append(Paragraph('APPLICATION/ANALYSIS', styles['Heading3']))
            story.append(Paragraph(irac_analysis['analysis'], styles['Normal']))
            story.append(Spacer(1, 12))

        if irac_analysis.get('conclusion'):
            story.append(Paragraph('CONCLUSION', styles['Heading3']))
            story.append(Paragraph(irac_analysis['conclusion'], styles['Normal']))
            story.append(Spacer(1, 12))

    async def _generate_prayer_for_relief_pdf(self, story, content, title_style, normal_style):
        """Generate Prayer for Relief PDF"""

        story.append(Paragraph(content['title'], title_style))
        story.append(Spacer(1, 20))

        story.append(Paragraph('WHEREFORE, Plaintiff prays for judgment as follows:', styles['Heading2']))
        story.append(Spacer(1, 12))

        remedies = content.get('remedies', [])
        for i, remedy in enumerate(remedies, 1):
            story.append(Paragraph(f"{i}. {remedy}", normal_style))
            story.append(Spacer(1, 12))

        story.append(Spacer(1, 20))
        story.append(Paragraph("Respectfully submitted,", normal_style))

    async def _generate_remedies_pdf(self, story, content, title_style, normal_style):
        """Generate Remedies PDF"""

        story.append(Paragraph(content['title'], title_style))
        story.append(Spacer(1, 20))

        damages = content.get('damages', {})
        if isinstance(damages, dict):
            for damage_type, amount in damages.items():
                story.append(Paragraph(f"{damage_type}: {amount}", normal_style))
                story.append(Spacer(1, 12))

    async def _generate_generic_pdf(self, story, content, title_style, normal_style):
        """Generate a generic PDF document"""

        story.append(Paragraph(content.get('title', 'Legal Document'), title_style))
        story.append(Spacer(1, 20))

        # Add content sections
        for key, value in content.items():
            if key != 'title' and isinstance(value, (str, list)):
                story.append(Paragraph(key.replace('_', ' ').title(), styles['Heading2']))
                story.append(Spacer(1, 12))
                if isinstance(value, list):
                    for item in value:
                        story.append(Paragraph(f"• {item}", normal_style))
                        story.append(Spacer(1, 6))
                else:
                    story.append(Paragraph(value, normal_style))
                story.append(Spacer(1, 12))

    async def _generate_placeholder(self, document_type: str, extension: str) -> str:
        """Generate a placeholder file when dependencies are missing"""

        filename = f"{document_type}_placeholder_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{extension}"
        file_path = self.output_dir / filename

        with open(file_path, 'w') as f:
            f.write(f"Placeholder {extension.upper()} document: {document_type}\n")
            f.write("Generated by LawyerFactory AI Maestro\n")
            f.write(f"Timestamp: {datetime.now().isoformat()}\n")

        return str(file_path)


class MarkdownDocumentGenerator:
    """Generates Markdown documents for UI display and collaboration"""

    def __init__(self):
        self.output_dir = Path("output/orchestration/markdown")
        self.output_dir.mkdir(exist_ok=True, parents=True)

    async def generate(self, document_type: str, content: Dict[str, Any], metadata: Dict[str, Any]) -> str:
        """Generate a Markdown document"""

        filename = f"{document_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
        file_path = self.output_dir / filename

        markdown_content = await self._generate_markdown_content(document_type, content, metadata)

        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(markdown_content)

        return str(file_path)

    async def _generate_markdown_content(self, document_type: str, content: Dict[str, Any], metadata: Dict[str, Any]) -> str:
        """Generate Markdown content"""

        md = []

        # Header
        md.append(f"# {content.get('title', 'Legal Document')}")
        md.append("")

        # Metadata
        md.append("## Document Information")
        md.append(f"- **Generated by:** LawyerFactory AI Maestro")
        md.append(f"- **Timestamp:** {datetime.now().isoformat()}")
        md.append(f"- **Case:** {metadata.get('case_name', 'Unknown Case')}")
        md.append(f"- **Jurisdiction:** {metadata.get('jurisdiction', 'Unknown')}")
        md.append("")

        # Generate content based on type
        if document_type == 'statement_of_facts':
            md.extend(await self._generate_statement_of_facts_md(content))
        elif document_type == 'claims_of_action':
            md.extend(await self._generate_claims_of_action_md(content))
        elif document_type == 'prayer_for_relief':
            md.extend(await self._generate_prayer_for_relief_md(content))
        elif document_type == 'remedies':
            md.extend(await self._generate_remedies_md(content))
        else:
            md.extend(await self._generate_generic_md(content))

        return "\n".join(md)

    async def _generate_statement_of_facts_md(self, content: Dict[str, Any]) -> List[str]:
        """Generate Statement of Facts Markdown"""

        md = []
        md.append("## Statement of Facts")
        md.append("")

        intro_text = f"*This action is brought in the {content['jurisdiction']} {content['court']}.*"
        md.append(intro_text)
        md.append("")

        facts = content.get('facts', [])
        for i, fact in enumerate(facts, 1):
            md.append(f"{i}. {fact}")
        md.append("")

        return md

    async def _generate_claims_of_action_md(self, content: Dict[str, Any]) -> List[str]:
        """Generate Claims of Action Markdown with IRAC"""

        md = []
        md.append("## Claims of Action")
        md.append("")

        claims = content.get('claims', [])
        for i, claim in enumerate(claims, 1):
            if isinstance(claim, dict):
                claim_title = claim.get('title', f'Claim {i}')
                md.append(f"### {i}. {claim_title}")
                md.append("")

                if content.get('use_irac', False):
                    irac_analysis = claim.get('irac_analysis', {})
                    md.extend(await self._add_irac_md(irac_analysis))
                else:
                    elements = claim.get('elements', [])
                    for element in elements:
                        md.append(f"- {element}")

                md.append("")

        return md

    async def _add_irac_md(self, irac_analysis: Dict[str, Any]) -> List[str]:
        """Add IRAC methodology to Markdown"""

        md = []

        if irac_analysis.get('issue'):
            md.append("#### ISSUE")
            md.append(irac_analysis['issue'])
            md.append("")

        if irac_analysis.get('rule'):
            md.append("#### RULE")
            md.append(irac_analysis['rule'])
            md.append("")

        if irac_analysis.get('analysis'):
            md.append("#### APPLICATION/ANALYSIS")
            md.append(irac_analysis['analysis'])
            md.append("")

        if irac_analysis.get('counterarguments'):
            md.append("#### COUNTERARGUMENTS")
            counterargs = irac_analysis['counterarguments']
            for counter in counterargs:
                md.append(f"- {counter}")
            md.append("")

        if irac_analysis.get('conclusion'):
            md.append("#### CONCLUSION")
            md.append(irac_analysis['conclusion'])
            md.append("")

        return md

    async def _generate_prayer_for_relief_md(self, content: Dict[str, Any]) -> List[str]:
        """Generate Prayer for Relief Markdown"""

        md = []
        md.append("## Prayer for Relief")
        md.append("")

        md.append("**WHEREFORE, Plaintiff prays for judgment as follows:**")
        md.append("")

        remedies = content.get('remedies', [])
        for i, remedy in enumerate(remedies, 1):
            md.append(f"{i}. {remedy}")
        md.append("")

        md.append("*Respectfully submitted,*")
        md.append("")

        return md

    async def _generate_remedies_md(self, content: Dict[str, Any]) -> List[str]:
        """Generate Remedies Markdown"""

        md = []
        md.append("## Remedies")
        md.append("")

        damages = content.get('damages', {})
        if isinstance(damages, dict):
            for damage_type, amount in damages.items():
                md.append(f"- **{damage_type}:** {amount}")
        md.append("")

        return md

    async def _generate_generic_md(self, content: Dict[str, Any]) -> List[str]:
        """Generate a generic Markdown document"""

        md = []
        md.append("## Content")
        md.append("")

        # Add content sections
        for key, value in content.items():
            if key != 'title' and isinstance(value, (str, list)):
                md.append(f"### {key.replace('_', ' ').title()}")
                md.append("")
                if isinstance(value, list):
                    for item in value:
                        md.append(f"- {item}")
                else:
                    md.append(value)
                md.append("")

        return md


# Global instance
document_generator = LegalDocumentGenerator()